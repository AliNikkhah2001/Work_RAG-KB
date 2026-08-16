"""DOCX parser using python-docx for structured text extraction."""

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from kb_manager.parsers.base import BaseParser, ParsedDocument


# Common heading styles in DOCX
_HEADING_STYLE_PREFIXES = ("heading", "title", "subtitle")

# Heading level pattern (e.g., "Heading 1", "Heading 2")
_HEADING_LEVEL_PATTERN = re.compile(r"heading\s*(\d+)", re.IGNORECASE)


def _detect_heading_level(paragraph: object) -> int | None:
    """Detect if a paragraph is a heading and return its level.

    Args:
        paragraph: A python-docx paragraph object.

    Returns:
        Heading level (1-6) or None if not a heading.
    """
    style_name = (paragraph.style.name or "").strip()
    style_lower = style_name.lower()

    # Check style name prefix
    for prefix in _HEADING_STYLE_PREFIXES:
        if style_lower.startswith(prefix):
            match = _HEADING_LEVEL_PATTERN.search(style_name)
            if match:
                return int(match.group(1))
            # Default levels for title/subtitle
            if prefix == "title":
                return 1
            if prefix == "subtitle":
                return 2
            return 1

    # Check for inline font-based headings (bold + larger font)
    if paragraph.runs:
        run = paragraph.runs[0]
        if run.bold and run.font.size:
            point_size = run.font.size.pt
            if point_size >= 18:
                return 1
            if point_size >= 14:
                return 2
            if point_size >= 12:
                return 3

    return None


def _extract_paragraph_text(paragraph: object) -> str:
    """Extract text from a paragraph, preserving formatting markers.

    Args:
        paragraph: A python-docx paragraph object.

    Returns:
        Extracted text with bold/italic markers.
    """
    parts: list[str] = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text.strip():
            parts.append(text)
            continue

        # Add formatting markers for content extraction
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"

        parts.append(text)

    return "".join(parts)


def _build_content_title(doc: object, file_path: str) -> str:
    """Derive document title from metadata or filename.

    Args:
        doc: python-docx Document object.
        file_path: Original file path.

    Returns:
        Document title string.
    """
    # Try core properties first
    if hasattr(doc, "core_properties"):
        props = doc.core_properties
        if props.title and props.title.strip():
            return props.title.strip()
        if props.subject and props.subject.strip():
            return props.subject.strip()

    return Path(file_path).stem


class DocxParser(BaseParser):
    """Parser for Microsoft Word DOCX files.

    Extracts paragraphs with heading detection and basic formatting
    preservation (bold, italic markers).
    """

    def can_parse(self, file_path: str) -> bool:
        """Check if the file is a DOCX that can be parsed.

        Args:
            file_path: Path to check.

        Returns:
            True if the file has a .docx extension.
        """
        return Path(file_path).suffix.lower() == ".docx"

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a DOCX file and extract structured content.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            ParsedDocument with content and detected sections.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file cannot be opened.
        """
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        abs_path = str(Path(file_path).resolve())

        try:
            doc = Document(abs_path)
        except Exception as e:
            raise ValueError(f"Failed to open DOCX file: {e}") from e

        title = _build_content_title(doc, file_path)

        metadata: dict = {}
        if hasattr(doc, "core_properties"):
            props = doc.core_properties
            metadata = {
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "paragraph_count": len(doc.paragraphs),
            }

        # Extract all paragraphs and build sections
        all_text_parts: list[str] = []
        sections: list[dict] = []
        current_section: dict | None = None

        for para in doc.paragraphs:
            text = _extract_paragraph_text(para)
            if not text.strip():
                continue

            heading_level = _detect_heading_level(para)

            if heading_level is not None:
                # Start a new section
                if current_section is not None:
                    sections.append(current_section)
                current_section = {
                    "heading": text.strip(),
                    "text": "",
                    "level": heading_level,
                }
            elif current_section is not None:
                # Append to current section
                if current_section["text"]:
                    current_section["text"] += "\n"
                current_section["text"] += text
            else:
                # Text before any heading
                all_text_parts.append(text)

        # Don't forget the last section
        if current_section is not None:
            sections.append(current_section)

        # Combine all content
        content_parts = all_text_parts + [
            f"{'#' * s['level']} {s['heading']}\n{s['text']}" for s in sections
        ]
        combined_content = "\n\n".join(content_parts)

        if not combined_content.strip():
            raise ValueError(f"No text content extracted from DOCX: {file_path}")

        # Clean up sections (remove level key for output)
        clean_sections = [{"heading": s["heading"], "text": s["text"]} for s in sections]

        return ParsedDocument(
            source_path=abs_path,
            title=title,
            content=combined_content,
            file_type="docx",
            metadata=metadata,
            sheets=None,
            sections=clean_sections if clean_sections else None,
        )
