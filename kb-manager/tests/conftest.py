from __future__ import annotations

import hashlib
import random
from typing import TYPE_CHECKING

import pytest
import pytest_asyncio

if TYPE_CHECKING:
    from pathlib import Path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class MockEmbedder:
    """Return deterministic random vectors for testing."""

    def __init__(self, dimensions: int = 384):
        self._dimensions = dimensions
        self._cache: dict[str, list[float]] = {}

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        results: list[list[float]] = []
        for text in texts:
            h = hashlib.sha256(text.encode()).hexdigest()
            seed = int(h[:8], 16)
            rng = random.Random(seed)
            vec = [rng.random() for _ in range(self._dimensions)]
            results.append(vec)
            self._cache[h] = vec
        return results

    def embed_query(self, query: str) -> list[float]:
        return self.embed_texts([query])[0]

    @property
    def dimensions(self) -> int:
        return self._dimensions

    @property
    def model_name(self) -> str:
        return "mock-embedder"


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def tmp_source_dir(tmp_path: Path) -> Path:
    """Create a temporary directory with sample test files."""
    (tmp_path / "reason_codes.xlsx").touch()
    (tmp_path / "guide.docx").touch()
    (tmp_path / "notes.txt").write_text("some notes", encoding="utf-8")
    return tmp_path


@pytest.fixture()
def sample_xlsx(tmp_path: Path) -> Path:
    """Create a sample Excel file with reason code data."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "ReasonCodes"
    ws.append(["reason_code", "description", "category"])
    ws.append(["RC001", "خطای سرور", "سیستمی"])
    ws.append(["RC002", "خطای احراز هویت", "امنیتی"])
    ws.append(["RC003", "خطای شبکه", "شبکه"])

    path = tmp_path / "sample_reasons.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Create a sample DOCX file."""
    from docx import Document

    doc = Document()
    doc.add_heading("راهنمای سیستم", level=1)
    doc.add_paragraph("این یک سند آزمایشی است.")
    doc.add_heading("بخش اول", level=2)
    doc.add_paragraph("محتوای بخش اول.")

    path = tmp_path / "sample_guide.docx"
    doc.save(str(path))
    return path


@pytest_asyncio.fixture()
async def db_engine():
    """Create an in-memory SQLite async engine for unit testing."""
    from sqlalchemy import text
    from sqlalchemy.ext.asyncio import create_async_engine

    engine = create_async_engine("sqlite+aiosqlite:///:memory:")
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "CREATE TABLE IF NOT EXISTS documents ("
                "  id TEXT PRIMARY KEY,"
                "  source_path TEXT,"
                "  file_hash TEXT,"
                "  created_at TIMESTAMP"
                ")"
            )
        )
        await conn.execute(
            text(
                "CREATE TABLE IF NOT EXISTS chunks ("
                "  id TEXT PRIMARY KEY,"
                "  document_id TEXT,"
                "  content TEXT,"
                "  heading_path TEXT,"
                "  chunk_index INTEGER,"
                "  token_count INTEGER,"
                "  embedding BLOB"
                ")"
            )
        )
    yield engine
    await engine.dispose()


@pytest_asyncio.fixture()
async def db_session(db_engine):
    """Provide an async session fixture."""
    from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

    session_factory = async_sessionmaker(db_engine, class_=AsyncSession, expire_on_commit=False)
    async with session_factory() as session:
        yield session


@pytest.fixture()
def preprocessor():
    """Return a PersianPreprocessor instance."""
    from kb_manager.preprocessor import PersianPreprocessor

    return PersianPreprocessor(spell_check=False)


@pytest.fixture()
def preprocessor_pipeline():
    """Return a PreprocessingPipeline instance."""
    from kb_manager.preprocessor.pipeline import PreprocessingPipeline

    return PreprocessingPipeline(spell_check=False)


@pytest.fixture()
def chunker():
    """Return a SemanticChunker instance."""
    from kb_manager.chunker import SemanticChunker

    return SemanticChunker(max_tokens=512, min_tokens=100)


@pytest.fixture()
def fixed_chunker():
    """Return a FixedChunker instance."""
    from kb_manager.chunker import FixedChunker

    return FixedChunker(max_tokens=512, min_tokens=100)


@pytest.fixture()
def embedder() -> MockEmbedder:
    """Return a mock embedder that returns random vectors."""
    return MockEmbedder(dimensions=384)
